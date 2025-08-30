#!/usr/bin/env python3
"""
Meeting‑to‑DOCX pipeline - “chunked” edition + paragraph timestamps
July 2025 (rev F‑prompt‑selector)

Change log (rev F‑prompt‑selector)
==================================
• **Output format selector**: added CLI flag `--format {meeting|lecture|qa}`.
  ‑ `meeting` keeps the original executive‑summary prompt (default).
  ‑ `lecture` rewrites a talk/lecture into a clean, well‑structured text.
  ‑ `qa` extracts every question, its answer, and final suggestions from Q&A sessions.
• **Summariser refactor**: `summarise()` now takes `mode` and switches system/instruction prompts accordingly.
• **Minor**: docstrings, help strings, and variable names updated to reflect the new feature.

(Previous change log entries preserved below)
===============================
• **Updated folder structure**: the original media file is copied to
  `./data/audio_files/` for archival, and the generated DOCX transcript is
  saved to `./data/transcripts/`.
• **Summary model switched**: summarisation now defaults to `o4-mini`, per
  OpenAI’s docs: https://platform.openai.com/docs/models/o4-mini
• **Chat completion params updated**: using `max_completion_tokens` instead
  of `max_tokens` for `o4-mini` compatibility.
"""

from __future__ import annotations

import argparse
import shutil
import sys
import tempfile
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Type

from dotenv import load_dotenv
from docx import Document
from imageio_ffmpeg import get_ffmpeg_exe
from openai import OpenAI, APIConnectionError, InternalServerError
import httpx

import os                   
from dotenv import load_dotenv  
load_dotenv()

AUDIO_RATE = 16_000  # Hz
MAX_UPLOAD_MB = 10  # Hard‑limit per OpenAI docs (≈ 25 MB per request)
_BYTES_PER_SEC = AUDIO_RATE * 2  # 16‑bit mono → 2 bytes per sample
CHUNK_SECONDS = max(1, (MAX_UPLOAD_MB * 1024 ** 2) // _BYTES_PER_SEC)  # ≈ 787 s

# -----------------------------------------------------------------------------
# Models & formats
# -----------------------------------------------------------------------------
# -----------------------------------------------------------------------------  
# Models & formats  
# -----------------------------------------------------------------------------  

# Only the models that really return segment‑level timestamps stay here.
SEGMENT_MODELS: set[str] = {
    "whisper-1",
    "gpt-4o-mini-transcribe",
    }

# Models that *must* use plain "json" even though we might want segments.
_JSON_ONLY_MODELS: set[str] = {
    "gpt-4o-transcribe",          # CLI flag you pass
    "gpt-4o-transcribe-api-ev3",  # actual name seen by the API
}

def _supports_segments(model: str) -> bool:
    return model in SEGMENT_MODELS


def _response_format_for(model: str) -> str:
    # Force "json" for models that don’t understand "verbose_json".
    if model in _JSON_ONLY_MODELS:
        return "json"
    return "verbose_json" if _supports_segments(model) else "json"


# -----------------------------------------------------------------------------
# Utility
# -----------------------------------------------------------------------------

def format_timestamp(seconds: float) -> str:
    h, rem = divmod(int(seconds), 3600)
    m, s = divmod(rem, 60)
    return f"{h:02}:{m:02}:{s:02}"

# -----------------------------------------------------------------------------
# Paragraph grouping
# -----------------------------------------------------------------------------

def paragraphs_from_segments(
    segs: List[Dict[str, Any]],
    *,
    time_offset: float,
    min_chars: int = 300,
) -> List[str]:
    lines: List[str] = []
    buf: List[str] = []
    para_start: Optional[float] = None

    def flush():
        if buf and para_start is not None:
            ts = format_timestamp(para_start)
            lines.append(f"[{ts}] {' '.join(buf).strip()}")
            buf.clear()

    for seg in segs:
        if para_start is None:
            para_start = time_offset + seg["start"]
        buf.append(seg["text"].strip())
        ends_sentence = buf[-1].endswith((".", "!", "?"))
        if ends_sentence and sum(map(len, buf)) >= min_chars:
            flush()
            para_start = None

    flush()
    return lines

# -----------------------------------------------------------------------------
# Audio helpers
# -----------------------------------------------------------------------------
from subprocess import run

def extract_audio(src: Path, wav_out: Path) -> Path:
    ffmpeg = get_ffmpeg_exe()
    cmd = [
        ffmpeg,
        "-loglevel",
        "error",
        "-y",
        "-i",
        str(src),
        "-ar",
        str(AUDIO_RATE),
        "-ac",
        "1",
        str(wav_out),
    ]
    run(cmd, check=True, close_fds=sys.platform != "win32")
    return wav_out

def chunk_audio(wav: Path, chunk_dir: Path) -> List[Path]:
    limit_bytes = MAX_UPLOAD_MB * 1024 ** 2
    if wav.stat().st_size <= limit_bytes:
        return [wav]
    chunk_dir.mkdir(parents=True, exist_ok=True)
    ffmpeg = get_ffmpeg_exe()
    out_pattern = chunk_dir / "chunk_%04d.wav"
    cmd = [
        ffmpeg,
        "-loglevel",
        "error",
        "-y",
        "-i",
        str(wav),
        "-f",
        "segment",
        "-segment_time",
        str(CHUNK_SECONDS),
        "-c",
        "copy",
        str(out_pattern),
    ]
    run(cmd, check=True, close_fds=sys.platform != "win32")
    chunks = sorted(chunk_dir.glob("chunk_*.wav"))
    if not chunks:
        raise RuntimeError("Chunking failed: no output segments created.")
    return chunks

# -----------------------------------------------------------------------------
# Retry helper
# -----------------------------------------------------------------------------

def _with_retries(fn, *, max_retries: int, backoff: float = 2.0, exc_types: tuple[Type[BaseException], ...]):
    for attempt in range(1, max_retries + 1):
        try:
            return fn()
        except exc_types as e:
            if attempt == max_retries:
                raise
            wait = backoff ** (attempt - 1)
            print(f"⚠️ {e.__class__.__name__}: {e} - retry {attempt}/{max_retries - 1} in {wait:.1f}s …")
            time.sleep(wait)

# -----------------------------------------------------------------------------
# OpenAI STT
# -----------------------------------------------------------------------------

def transcribe_chunk(
    wav: Path,
    *,
    model: str,
    timeout_s: float,
    max_retries: int,
) -> Dict[str, Any]:
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"),
                    timeout=timeout_s)
    fmt = _response_format_for(model)
    print(f"⏳ Transcribing {wav.name} with {model} … (format: {fmt})")

    def _call():
        with wav.open("rb") as fh:
            return client.audio.transcriptions.create(
                model=model,
                file=fh,
                response_format=fmt,
                timestamp_granularities=["segment"] if fmt == "verbose_json" else None,
            )

    resp = _with_retries(
        _call, max_retries=max_retries, exc_types=(APIConnectionError, InternalServerError, httpx.ReadError)
    )
    data = resp.model_dump()
    return {"text": data.get("text", ""), "segments": data.get("segments", [])}

# -----------------------------------------------------------------------------
# Summarisation
# -----------------------------------------------------------------------------

def summarise(
    text: str,
    *,
    mode: str = "meeting",  # {meeting|lecture|qa}
    model: str ="o4-mini-2025-04-16",
    timeout_s: float = 60.0,
) -> str:
    """Generate a structured summary according to *mode*.

    • ``meeting`` → executive summary & action items (original behaviour)
    • ``lecture`` → cleaned, well‑written lecture/talk text
    • ``qa``      → list of Q&As plus suggestions
    """

    client = OpenAI(timeout=timeout_s)

    if mode == "meeting":
        system_msg = (
            "Agisci come un esperto facilitatore di riunioni e project‑manager. "
            "Il tuo compito è distillare trascrizioni in sintesi utilizzabili da dirigenti italiani. "
            "Il risultato deve essere neutro, preciso e professionale. "
            "Pensa passo‑passo internamente ma NON rivelare il ragionamento."
        )
        instructions = (
            "Produci SOLO il testo seguente, senza markdown:\n\n"
            "RIEPILOGO ESECUTIVO\n"
            "• 5‑7 bullet (≤ 20 parole) con verbi all'infinito.\n\n"
            "AZIONI E RESPONSABILI\n"
            "• Formato: <Responsabile> → <Azione> (scadenza)\n"
            "• Se la scadenza non è indicata, scrivi “data da definire”."
        )
    elif mode == "lecture":
        system_msg = (
            "Sei un Revisore Accademico e Redattore Professionista esperto nella trasformazione di trascrizioni orali in testi scritti di elevate qualità. "
            "Quando ricevi una trascrizione di lezione o talk, procedi così:\n"
            "1. (Hidden) Analizza passo‑passo la struttura e i contenuti con ragionamento interno per organizzare logicamente l’informazione.\n"
            "2. Struttura il testo in sezioni distinte con titoli descrittivi e paragrafi scorrevoli.\n"
            "Se riscontri ambiguità, contrassegna l’area con “[DA CHIARIRE]”."
            )

        instructions = (
            "Produci DUE sezioni distinte (senza includere il ragionamento interno):\n\n"
            "SEZIONE 1 - RIASSUNTO (max 100 parole)\n"
            "- Elenca in 5–7 bullet point i punti chiave della lezione.\n\n"
            "SEZIONE 2 - TESTO RIELABORATO\n"
            "Titolo: <Titolo conciso e descrittivo>\n"
            "Paragrafo 1: <Introduzione sintetica>\n"
            "Paragrafo 2: <Sviluppo strutturato con eventuali sottotitoli>\n"
            "…\n\n"
            "Indicazioni di stile: registro formale, coerenza terminologica, tempi verbali appropriati, paragrafi brevi."
            )
    elif mode == "qa":
        system_msg = (
            "Agisci come stenografo e facilitatore di sessioni Q&A. "
            "Devi estrarre ogni domanda, la relativa risposta solo se presente, e riepilogare i suggerimenti finali."
        )
        instructions = (
            "Analizza (nascosto) e poi produci SOLO il testo:\n\n"
            "DOMANDE E RISPOSTE\n"
            "Q: <domanda 1>\nA: <risposta 1 (se presente)>\n\n[… altre Q/A …]\n\n"
            "SUGGERIMENTI EMERSI\n"
            "- <suggerimento 1>\n- …"
        )
    else:
        raise ValueError(f"Unknown summary mode: {mode}")

    user_msg = f"transcript:\n {text}"
    temper = 1 if model == "o4-mini" else 0.0

    print("⏳ Generazione del riepilogo …")
    res = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": instructions},
            {"role": "user", "content": user_msg},
        ],
        temperature=temper,
        max_completion_tokens=10000,
    )
    return res.choices[0].message.content.strip()

# ──────────────────────────────────────────────────────────────────────────────
# DOCX writers
# ──────────────────────────────────────────────────────────────────────────────
def save_docx(transcript: str, summary: str | None, dst: Path) -> None:
    """Write transcript and (optionally) a summary into *dst*."""
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if summary:                                               # summary first
        doc.add_heading("Summary", level=1)
        for ln in summary.splitlines():
            doc.add_paragraph(ln)
        doc.add_page_break()
    doc.add_heading("Transcript", level=1)
    for ln in transcript.splitlines():
        if ln.strip():
            doc.add_paragraph(ln)
    doc.save(dst)

def summarise_docx_file(
docx_path: str, *, mode: str, summarize_model: str, timeout_s: float
) -> None:
    """Prefix *docx_path* with a fresh summary (in‑place)."""
    src = Path(docx_path)
    if not src.exists():
        sys.exit(f"❌ DOCX not found: {src}")
    tmp_out = src.with_suffix(".tmp.docx")

    orig = Document(src)
    full_text = "\n".join(p.text for p in orig.paragraphs)
    summary = summarise(full_text, mode=mode, model=summarize_model, timeout_s=timeout_s)

    new_doc = Document()
    new_doc.add_heading("Summary", level=1)
    for ln in summary.splitlines():
        new_doc.add_paragraph(ln)
    new_doc.add_page_break()
    # copy original text (formatting lost but minimal edit as requested)
    for p in orig.paragraphs:
        new_doc.add_paragraph(p.text)
    new_doc.save(tmp_out)
    shutil.move(tmp_out, src)        # overwrite atomically
    print(f"📄 Summary inserted: {src.absolute()}")

# ──────────────────────────────────────────────────────────────────────────────
# Orchestrator
# ──────────────────────────────────────────────────────────────────────────────
def transcribe_file(
input_path: str, *, stt_model: str, summarize_model: str, mode: str,
    timeout_s: float, max_retries: int, include_summary: bool
) -> None:
    load_dotenv()
    src = Path(input_path)
    if not src.exists():
        sys.exit(f"❌ File not found: {src}")

    transcripts_dir = Path("./data/transcripts")
    audio_dir = Path("./data/audio_files")
    transcripts_dir.mkdir(parents=True, exist_ok=True)
    audio_dir.mkdir(parents=True, exist_ok=True)

    dst_docx = transcripts_dir / src.with_suffix(".docx").name
    if src.suffix.lower() != ".wav":             # archive original media
        try:
            shutil.copy2(src, audio_dir / src.name)
        except Exception as e:
            print(f"⚠️  Unable to archive original media: {e}")

    with tempfile.TemporaryDirectory() as td:
        td_path = Path(td)
        wav_path = td_path / "full.wav"
        extract_audio(src, wav_path)
        chunks = chunk_audio(wav_path, td_path / "chunks")

        plain_parts, timestamped_lines = [], []
        time_offset = 0.0
        for idx, cp in enumerate(chunks, 1):
            print(f"— Chunk {idx}/{len(chunks)} ({cp.stat().st_size/1_048_576:.1f} MB)")
            stt = transcribe_chunk(cp, model=stt_model,
                                   timeout_s=timeout_s, max_retries=max_retries)
            plain_parts.append(stt["text"])
            segs = stt.get("segments", [])
            if segs:
                timestamped_lines += paragraphs_from_segments(segs, time_offset=time_offset)
                time_offset += segs[-1].get("end", 0.0)
            else:
                timestamped_lines.append(f"[{format_timestamp(time_offset)}] {stt['text'].strip()}")
                time_offset += cp.stat().st_size / _BYTES_PER_SEC

    transcript_txt = "\n".join(timestamped_lines)
    summary_txt = (
        summarise("\n".join(plain_parts), mode=mode, model=summarize_model, timeout_s=timeout_s)
        if include_summary else None
    )
    save_docx(transcript_txt, summary_txt, dst_docx)
    print(f"📄 Transcript saved: {dst_docx.absolute()}")

# ──────────────────────────────────────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────────────────────────────────────
def parse_args():
    p = argparse.ArgumentParser(description="Audio / video → DOCX (or DOCX → add summary)")
    p.add_argument("input_file", help="Input media or DOCX file")
    p.add_argument("--task", choices=["full", "transcript", "summarize"],
                   default="full",
                   help="full (default): summary+transcript | "
                        "transcript: transcript‑only | summarize: add summary to DOCX")
    p.add_argument("--model", default="whisper-1",
                   help="STT model (for audio tasks, default: whisper-1)")
    p.add_argument("--summarize-model", default="o4-mini",
                   help="Summarization model (default: o4-mini)")
    p.add_argument("--format", choices=["meeting", "lecture", "qa"],
                   default="meeting",
                   help="Summary style (meeting, lecture, qa)")
    p.add_argument("--timeout", type=float, default=60,
                   help="Per‑request timeout (s)")
    p.add_argument("--max-retries", type=int, default=4,
                   help="Retries on transient errors")
    return p.parse_args()

# ──────────────────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    args = parse_args()

    if args.task == "summarize":
        if Path(args.input_file).suffix.lower() != ".docx":
            sys.exit("❌ --task summarize expects a .docx file as input.")
        summarise_docx_file(
            args.input_file,
            mode=args.format,
            summarize_model=args.summarize_model,
            timeout_s=args.timeout
        )

    else:  # audio/video tasks
        include_summary = args.task == "full"
        transcribe_file(
            input_path=args.input_file,
            stt_model=args.model,
            summarize_model=args.summarize_model,
            mode=args.format,
            timeout_s=args.timeout,
            max_retries=args.max_retries,
            include_summary=include_summary,
        )